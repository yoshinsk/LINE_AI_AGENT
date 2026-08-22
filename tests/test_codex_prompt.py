r"""<PROJECT_ROOT>\tests\test_codex_prompt.py

Codexへ渡すプロンプトに会話履歴、検索ナレッジ、添付パスが入ることを検証します。
"""

from __future__ import annotations

import json
from pathlib import Path
import sys
import tempfile
import unittest

from docx import Document

from line_ai_agent.codex_runner import OFFICE_REVISION_OUTPUT_SCHEMA, CodexJob, CodexRunner, _tail, build_office_revision_prompt, build_prompt, requires_office_revision
from line_ai_agent.projects import ProjectSelection


class CodexPromptTest(unittest.TestCase):
    """LINEジョブをCodex向け文脈へ変換する処理を確認します。"""

    def test_prompt_contains_attachment_paths_and_knowledge(self) -> None:
        job = CodexJob(
            job_id=12,
            source_key="user:Uxxx",
            request_text="添付を要約してください。",
            project=ProjectSelection("none", None, None, "未指定"),
            recent_messages=({"role": "user", "body": "前回の相談", "created_at": "2026-08-20"},),
            knowledge=({"role": "assistant", "text": "過去の回答", "created_at": "2026-08-20"},),
            attachments=(Path("C:/tmp/report.pdf"),),
        )
        prompt = build_prompt(job)
        self.assertIn(str(Path("C:/tmp/report.pdf")), prompt)
        self.assertIn("過去の回答", prompt)
        self.assertIn("添付を要約してください。", prompt)

    def test_codex_command_attaches_image_files(self) -> None:
        runner = CodexRunner(
            command='codex exec --output-last-message {output_file} -',
            timeout_seconds=30,
            no_project_workdir=Path("C:/tmp"),
            reply_max_chars=4500,
            result_asset_output_dir=Path("C:/tmp/result-assets"),
            result_asset_allowed_dirs=(Path("C:/tmp/result-assets"),),
            result_asset_max_count=5,
        )
        image_path = Path("C:/tmp/photo.jpg")
        pdf_path = Path("C:/tmp/report.pdf")
        job = CodexJob(
            job_id=34,
            source_key="group:Gxxx",
            request_text="画像を確認してください。",
            project=ProjectSelection("none", None, None, "未指定"),
            recent_messages=(),
            knowledge=(),
            attachments=(image_path, pdf_path),
        )
        args, output_file = runner._prepare_command(job)
        image_arg_index = args.index("--image")
        self.assertEqual(str(image_path.resolve(strict=False)), args[image_arg_index + 1])
        self.assertLess(image_arg_index, args.index("-"))
        self.assertNotIn(str(pdf_path.resolve(strict=False)), args)
        self.assertIsNotNone(output_file)

    def test_prompt_embeds_office_text_snapshot(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            original = root / "motivation.docx"
            snapshot = root / "7-motivation.line-office-extracted.txt"
            snapshot.write_text("本文: 交換留学を志望する理由です。", encoding="utf-8")
            job = CodexJob(
                job_id=35,
                source_key="user:Uxxx",
                request_text="添削してください。",
                project=ProjectSelection("none", None, None, "未指定"),
                recent_messages=(),
                knowledge=(),
                attachments=(original, snapshot),
            )

            prompt = build_prompt(job)

            self.assertIn(str(original), prompt)
            self.assertIn("本文: 交換留学を志望する理由です。", prompt)
            self.assertIn("元ファイルを読めないことだけを理由に回答を保留しないでください。", prompt)

    def test_office_revision_prompt_requires_a_real_returned_file(self) -> None:
        job = CodexJob(
            job_id=36,
            source_key="user:Uxxx",
            request_text="添付の文章を添削して修正してください。",
            project=ProjectSelection("none", None, None, "未指定"),
            recent_messages=(),
            knowledge=(),
            attachments=(Path("C:/tmp/motivation.docx"),),
            result_asset_dir=Path("C:/tmp/result-assets/job-36"),
        )

        prompt = build_prompt(job)
        revision_prompt = build_office_revision_prompt(job)

        self.assertTrue(requires_office_revision(job))
        self.assertIn("本文の提案だけで完了してはいけません。", prompt)
        self.assertIn("-revised.docx", prompt)
        self.assertIn("JSON Schemaに厳密に従う編集計画", revision_prompt)
        self.assertIn("kind=word_text", revision_prompt)
        self.assertIn("全項目を必ず含めてください", revision_prompt)

    def test_office_summary_request_does_not_require_a_revision_file(self) -> None:
        job = CodexJob(
            job_id=37,
            source_key="user:Uxxx",
            request_text="添付の内容を要約してください。",
            project=ProjectSelection("none", None, None, "未指定"),
            recent_messages=(),
            knowledge=(),
            attachments=(Path("C:/tmp/report.xlsx"),),
        )

        self.assertFalse(requires_office_revision(job))

    def test_structured_office_plan_becomes_a_returned_docx_file(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            source = root / "source.docx"
            document = Document()
            document.add_paragraph("旧表記")
            document.save(source)
            script = root / "plan_writer.py"
            plan = {
                "summary": "修正済みWord文書を作成しました。",
                "files": [
                    {
                        "source_file": source.name,
                        "edits": [{"kind": "word_text", "find": "旧表記", "replacement": "新表記"}],
                    }
                ],
            }
            script.write_text(
                "import json, sys\nfrom pathlib import Path\n"
                f"Path(sys.argv[1]).write_text({json.dumps(json.dumps(plan, ensure_ascii=False))}, encoding='utf-8')\n",
                encoding="utf-8",
            )
            runner = CodexRunner(
                command=f"{sys.executable} {script} {{output_file}}",
                timeout_seconds=30,
                no_project_workdir=root,
                reply_max_chars=4500,
                result_asset_output_dir=root / "result-assets",
                result_asset_allowed_dirs=(root / "result-assets",),
                result_asset_max_count=5,
            )
            job = CodexJob(
                job_id=38,
                source_key="user:Uxxx",
                request_text="添付を添削して修正してください。",
                project=ProjectSelection("none", None, None, "未指定"),
                recent_messages=(),
                knowledge=(),
                attachments=(source,),
            )

            result = runner.run_office_revision(job)

            self.assertTrue(result.ok, result.text)
            self.assertEqual("修正済みWord文書を作成しました。", result.text)
            self.assertEqual(1, len(result.asset_paths))
            self.assertEqual("新表記", Document(result.asset_paths[0]).paragraphs[0].text)

    def test_tail_keeps_the_last_codex_error_after_startup_warnings(self) -> None:
        text = "起動警告\n" * 400 + "ERROR: schema is invalid"

        shortened = _tail(text, 120)

        self.assertIn("ERROR: schema is invalid", shortened)
        self.assertTrue(shortened.startswith("...（起動時出力を省略）"))

    def test_office_revision_schema_requires_every_edit_property(self) -> None:
        item = OFFICE_REVISION_OUTPUT_SCHEMA["properties"]["files"]["items"]["properties"]["edits"]["items"]

        self.assertEqual(set(item["properties"]), set(item["required"]))


if __name__ == "__main__":
    unittest.main()
